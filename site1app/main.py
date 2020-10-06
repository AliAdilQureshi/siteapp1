# -------------first section is for crunching numbers and uploading these crunched numbers to a DB table------------

import pandas as pd
from site1app.models import FileUpload, saf_defect_table, Report_File_Upload_Table

# checking to see if there is a file in FileUpload DB Table that is logges as pftest.csv in the attribution file section of the table. then setting the pandas df to that file.
if FileUpload.objects.filter(attribution_file_upload='pftest4.csv').exists():
    df = pd.read_csv(FileUpload.objects.get(attribution_file_upload='pftest4.csv').attribution_file_upload)


# counting the various values in the status column and setting it to totals
totals = df["status"].value_counts()

# we know that the column contains the values ok and fail so we set variables to them.
oks = float(totals.OK)
fails = float(totals.Fail)

# next lines are just for if the db table saf defetc table is empty with no values, then we set the values to zero
if saf_defect_table.objects.count() == 0:
    saf_defect_table.objects.create(saf_number_in_table=oks, defect_number_in_table=fails)
else:
    crunched_numbers_for_upload = saf_defect_table.objects.all()  # if the values are not empty we will up date the values with the values we counted earlrier
    crunched_numbers_for_upload.update(saf_number_in_table=oks, defect_number_in_table=fails)

print(df)
print(oks)
print(fails)

# next step is to build a chart that can read the "crunched numbers" from the saf_defect_table
# ---------------end of section 1------------

# ----------------section 2 is for the generation of a pdf report file---------------

import docx
import io
from docx2pdf import convert
from django.http import FileResponse
import matplotlib.pyplot as plt
from docx.shared import Inches
import numpy as np
import requests
import json


def making_a_doc_function(request):
    # plot from uploaded data now need to make chart from live data that has been uploaded to the site. this will be done by either doing pandas crunching all over again to get the numbers that you want or by reading numbers tht have already been put in one of the tables. most ikley a combination of both

    doc = docx.Document()
    pk = 3
    print(3)
    if not saf_defect_table.objects.filter(pk=pk).exists():  # this accesses the saf_defect_table
        doc.add_heading("no text")
        doc.save('thisisdoc.docx')
    else:
        saf_defect_table_variable = saf_defect_table.objects.get(pk=pk)

        doc.add_heading(str(saf_defect_table_variable.saf_number_in_table))
        doc.add_heading(str(saf_defect_table_variable.defect_number_in_table))

        # making a plot from a csv file that was uploaded to DB table
        if FileUpload.objects.filter(attribution_file_upload='pftest4.csv').exists():
            life3 = pd.read_csv(FileUpload.objects.get(attribution_file_upload='pftest4.csv').attribution_file_upload)
            life3.plot(kind='bar', x='status', y='item')
            plt.title('my plot title')
            plt.xlabel('years')
            plt.ylabel('Age')
            memfile = io.BytesIO()  # this is needed in order to save the file to a temporary memory section
            plt.savefig(memfile)
            doc.add_picture(memfile, width=Inches(4))
            memfile.close()
            plt.clf()  # this shuts down the plot so then you can make a new.

        # making a plot out of numbers that are already stored in saf defect table DB table
        y1 = saf_defect_table_variable.saf_number_in_table
        y2 = saf_defect_table_variable.defect_number_in_table
        korea_scores = (y1, y2)
        col_count = 2
        bar_width = .2
        index = np.arange(col_count)
        plt.barh(index, korea_scores, bar_width, alpha=.4, label="Korea")
        memfile2 = io.BytesIO()
        plt.savefig(memfile2)
        doc.add_picture(memfile2, width=Inches(4))
        memfile2.close()
        print("heeeeeeeeeee")

    doc.save('thisisdoc.docx')


    url = 'https://pdf.to/v1/api'
    files = {'file': open('thisisdoc.docx', 'rb')}
    params = {'convert_to': 'pdf'}
    data = {'data': json.dumps(params)}
    headers = {'Authorization': 'ak_04f39680fbab46ecbbfe7e0b554263ff'}

    r = requests.post(
        url,
        files=files,
        data=data,
        headers=headers
    )

    with open('result.pdf', 'wb') as f:
        f.write(r.content)

    pdf = open('result.pdf', 'rb')

    response = FileResponse(pdf)
    
    return response
# return render(request, 'doc.html', response)

# ----------------end of section 2--------------------------------------------


import docx2pdf
from django.shortcuts import render
from docx import Document

def making_a_doc_with_doc_function(request):
# letters = string.ascii_lowercase
#     result_str = ''.join(random.choice(letters) for i in range(length))
#     print("Random string of length", length, "is:", result_str)
#plot from uploaded data now need to make chart from live data that has been uploaded to the site. this will be done by either doing pandas crunching all over again to get the numbers that you want or by reading numbers tht have already been put in one of the tables. most ikley a combination of both

    doc = docx.Document()
    pk = 3
    if not saf_defect_table.objects.filter(pk=pk).exists(): #this accesses the saf_defect_table
        doc.add_heading("no text")
        doc.save('thisisdoc.docx')
    else:
        saf_defect_table_variable = saf_defect_table.objects.get(pk=pk)

        doc.add_heading(str(saf_defect_table_variable.saf_number_in_table))
        doc.add_heading(str(saf_defect_table_variable.defect_number_in_table))

        #making a plot from a csv file that was uploaded to DB table
        if FileUpload.objects.filter(attribution_file_upload='pftest4.csv').exists():
            life3 = pd.read_csv(FileUpload.objects.get(attribution_file_upload='pftest4.csv').attribution_file_upload)
            life3.plot(kind='bar', x='status', y='item')
            plt.title('my plot title')
            plt.xlabel('years')
            plt.ylabel('Age')
            memfile = io.BytesIO() #this is needed in order to save the file to a temporary memory section
            plt.savefig(memfile)
            doc.add_picture(memfile, width=Inches(4))
            memfile.close()
            plt.clf()  #this shuts down the plot so then you can make a new.

        #making a plot out of numbers that are already stored in saf defect table DB table
        y1 = saf_defect_table_variable.saf_number_in_table
        y2 = saf_defect_table_variable.defect_number_in_table
        korea_scores = (y1, y2)
        col_count = 2
        bar_width = .2
        index = np.arange(col_count)
        plt.barh(index, korea_scores, bar_width, alpha=.4, label="Korea")
        memfile2 = io.BytesIO()
        plt.savefig(memfile2)
        doc.add_picture(memfile2, width=Inches(4))
        memfile2.close()


    doc.save('thisisdoc.docx')


    #converting the generated docx into a pdf file
    pdf = open('thisisdoc.docx', 'rb')
    response = FileResponse(pdf)

   # Report_File_Upload_Table.objects.create(generated_report_file_upload=output.pdf)
    return response
